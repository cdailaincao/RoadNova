from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "RoadNova_Research_Paper.docx"


def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, before=0, after=6, line=1.0, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def set_two_columns(section):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")


def set_one_column(section):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols[0].set(qn("w:num"), "1")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, after=0)
    r = p.add_run(text)
    set_font(r, 9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph(p, before=8 if level == 1 else 4, after=3)
    r = p.add_run(text)
    set_font(r, 10, bold=True)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, after=5, line=1.02)
    r = p.add_run(text)
    set_font(r, 10)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style=None)
    set_paragraph(p, after=3, line=1.0)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    r = p.add_run("- " + text)
    set_font(r, 10)
    return p


def add_reference(doc, idx, text):
    p = doc.add_paragraph()
    set_paragraph(p, after=2, line=1.0)
    r = p.add_run(f"[{idx}] {text}")
    set_font(r, 8)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.68)
section.right_margin = Inches(0.68)
set_one_column(section)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
styles["Normal"].font.size = Pt(10)

p = doc.add_paragraph()
set_paragraph(p, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("RoadNova: A Django-Based India Road Trip Planning System with Route-Aware Itineraries, Open Maps, Weather, Safety and Shared Expense Support")
set_font(r, 18, bold=True)

authors = [
    ("Fahad Ahmed", "fahadahmed.2k6@gmail.com"),
    ("Deepshika S Reddy", "deepshikasreddy@gmail.com"),
    ("Bhutada Richa Sujan", "richadarak2104@gmail.com"),
    ("Deepankar Anand", "aidoneus46@gmail.com"),
]
table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True
for i, (name, email) in enumerate(authors):
    cell = table.cell(i // 2, i % 2)
    p = cell.paragraphs[0]
    set_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(name + "\n")
    set_font(r, 10, bold=True)
    r = p.add_run("Department of Computer Science and Engineering\n")
    set_font(r, 9)
    r = p.add_run("Dayananda Sagar Academy of Technology and Management\n")
    set_font(r, 9)
    r = p.add_run("Bangalore, India\n")
    set_font(r, 9)
    r = p.add_run(email)
    set_font(r, 9)

p = doc.add_paragraph()
set_paragraph(p, before=8, after=4)
r = p.add_run("Abstract—")
set_font(r, 10, bold=True)
r = p.add_run(
    "Road trips in India require planning across long inter-city distances, variable terrain, fuel and toll budgets, live weather, emergency access and shared group costs. "
    "RoadNova is a Django-based web system that converts a user's origin, destination, dates, group size and interests into a route-aware day-wise itinerary for ten Indian cities: Bangalore, Goa, Manali, Delhi, Mumbai, Hyderabad, Munnar, Surat, Jaipur and Chennai. "
    "Unlike destination-only itinerary generators, RoadNova emphasizes the journey itself by splitting long routes into feasible daily driving segments, enforcing a 300 km/day safety threshold, inserting named halt cities, fuel stops, scenic breaks, food stops and stays, and refusing under-planned long trips until the user increases the trip duration. "
    "The implementation integrates Leaflet, OpenStreetMap tiles, OSRM road routing, Open-Meteo weather forecasts, curated Indian points of interest, cost estimation for fuel, tolls, food and hotels, an SOS panel and a group expense split module. "
    "This paper presents the system design, algorithmic workflow and implementation choices behind RoadNova, and positions the project against prior work in tourist trip design, open-source road routing, weather-aware route planning and shared travel cost coordination."
)
set_font(r, 10)

p = doc.add_paragraph()
set_paragraph(p, after=8)
r = p.add_run("Index Terms—")
set_font(r, 10, bold=True, italic=True)
r = p.add_run("road trip planner, Django, OpenStreetMap, OSRM, itinerary generation, Open-Meteo, SOS, group expense split")
set_font(r, 10, italic=True)

new_section = doc.add_section(WD_SECTION.CONTINUOUS)
set_two_columns(new_section)

heading(doc, "I. INTRODUCTION")
body(doc, "Tour planning research commonly treats itinerary construction as a constrained selection and scheduling problem, where users have limited time and must choose among points of interest, route costs and preference constraints [1], [2]. In the Indian road-trip setting, the problem becomes more practical and safety-sensitive: a route may span several states, highways and climates, and a naive itinerary can easily propose unrealistic single-day driving distances.")
body(doc, "RoadNova addresses this gap as a student-built, full-stack Django project focused on India-based road travel. The system is not only a list of hotels and attractions. It plans the actual travel sequence, including distance-controlled daily route segments, intermediate halt cities, fuel guidance, scenic breaks, stay recommendations, weather for every major stop and shared cost estimates.")
body(doc, "The current RoadNova prototype supports ten cities across India and generates route plans for any origin-destination pair among them. For nearby trips the system produces direct, detailed plans. For longer journeys, such as Bangalore to Delhi or Bangalore to Manali, it applies a maximum daily travel limit and asks the user to extend the date range if the trip cannot be completed safely within the selected days.")

heading(doc, "II. RELATED WORK")
body(doc, "Tourist Trip Design Problem (TTDP) studies provide the theoretical base for planning activities under time, budget and preference constraints. Vansteenwegen and Souffriau describe trip planning functions using Operations Research models, especially the Orienteering Problem and extensions [1]. Souffriau, Vansteenwegen and Van Oudheusden model personalized tourist guides by selecting attractions under user and time constraints [2]. More recent work extends TTDP to multimodal tourist planning with road and pedestrian networks [3].")
body(doc, "For map and route computation, RoadNova is aligned with open geospatial systems. Haklay and Weber present OpenStreetMap as a user-generated street map platform [4]. OSRM and OSRM-based interfaces demonstrate that OpenStreetMap road data can support efficient distance, duration and route geometry computation [5], [6].")
body(doc, "Weather has a measurable role in travel planning and transport safety. Lam, Shao and Sumalee model impacts of adverse weather on road networks under demand and supply uncertainty [7], while Zhang, Hu and Liao study dynamic path optimization during adverse weather [8]. These works motivate RoadNova's use of live city and halt-level forecasts rather than destination-only weather.")
body(doc, "Shared cost coordination is another important group-travel requirement. Recent research on ride cost sharing and debt settlement formalizes how common mobility costs can be allocated or settled efficiently among participants [9], [10]. RoadNova applies this idea in a practical trip dashboard by estimating total and per-person costs and enabling users to add shared expenses.")

heading(doc, "III. SYSTEM OBJECTIVES")
body(doc, "The design objectives of RoadNova are practical and user-facing. The system should generate trip plans that are believable for Indian road travel, visually clear for a browser user, and built completely inside Django.")
bullet(doc, "Generate date-aware itineraries from selected origin, destination, start date and end date.")
bullet(doc, "Limit daily route distance to approximately 300 km and warn users when the chosen dates are insufficient.")
bullet(doc, "Use named intermediate halt cities instead of generic placeholders.")
bullet(doc, "Display an open-source road map route rather than a straight-line connection.")
bullet(doc, "Estimate fuel, toll, food, hotel, total and per-person costs.")
bullet(doc, "Fetch live weather for both major cities and intermediate halt points.")
bullet(doc, "Provide emergency calling and India emergency numbers for police, ambulance and fire support.")
bullet(doc, "Support group expense splitting during or after the trip.")

heading(doc, "IV. ARCHITECTURE AND IMPLEMENTATION")
body(doc, "RoadNova uses Django as the backend web framework. The trips application contains route data, itinerary generation services, API views, database models, templates, static JavaScript and CSS. Django templates render the home and feature screens, while JavaScript handles panel navigation, trip-plan requests, map rendering, weather loading, point-of-interest queries and expense split updates.")
body(doc, "The backend exposes endpoints for trip generation, city weather, coordinate-based weather, OSM point-of-interest search and trip details. The frontend uses Leaflet for map interaction, OpenStreetMap map tiles for base maps, OSRM for road-route geometry and Open-Meteo for weather forecasts.")

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for cell, text in zip(hdr, ["Layer", "Technology", "Role in RoadNova"]):
    shade_cell(cell, "D9EAF7")
    set_cell_text(cell, text, True)
rows = [
    ("Backend", "Django", "Routes, templates, APIs, admin, data models and itinerary services"),
    ("Map", "Leaflet + OpenStreetMap", "Interactive open-source map display"),
    ("Routing", "OSRM", "Road geometry and route distance estimation"),
    ("Weather", "Open-Meteo", "Forecasts for cities and halt coordinates"),
    ("Frontend", "HTML, CSS, JavaScript", "Feature dashboard, itinerary cards, costs, SOS and split UI"),
]
for row in rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, row):
        set_cell_text(cell, text)

heading(doc, "V. ITINERARY GENERATION METHOD")
body(doc, "The itinerary generator first validates the user input and calculates the inclusive number of travel days from the selected start and end dates. It estimates road travel demand and compares the required distance against the 300 km/day threshold. If the trip duration is too short, the system returns a clear error explaining the approximate route length and the minimum number of days required.")
body(doc, "For valid plans, RoadNova selects intermediate halt cities from a curated Indian halt dataset. The halt selection is ordered geographically between origin and destination and optimized so that no daily segment is overloaded. OSRM distances are then applied to route legs where available, replacing generic equal-distance splits with realistic road-distance values. Each day contains a travel-oriented sequence: departure, fuel stop, scenic halt, food break and arrival at the daily stay city.")
body(doc, "The generated itinerary includes hotels, restaurants and categories such as temples, beaches, heritage, hill-station scenery and adventure depending on user-selected interests. This balances journey planning with destination exploration, while keeping the primary focus on safe road movement.")

heading(doc, "VI. COST, WEATHER AND SAFETY MODULES")
body(doc, "The cost module estimates fuel using user mileage, tolls as a route-level estimate, and daily food and hotel costs based on travelers and duration. The final output includes total cost and per-person split, which is also reused by the group expense split module.")
body(doc, "Weather is fetched through a coordinate-based API path so that forecasts are not limited to the ten base cities. For every generated route city or halt point, the frontend requests weather using the stop name, latitude, longitude and required forecast horizon. This gives travelers visibility into changing conditions along long trips.")
body(doc, "The SOS feature contains a one-tap call action and India emergency numbers: 112 for national emergency support, 100 for police, 108 for ambulance and 101 for fire brigade. The module is designed as a quick-access safety surface inside the same dashboard as the trip plan.")

heading(doc, "VII. USER INTERFACE DESIGN")
body(doc, "The final RoadNova interface opens with a simplified home page instead of showing all tools at once. Feature buttons lead to Planner, Map, Costs, Weather, SOS and Group Expense Split sections. A top dashboard keeps the same feature areas accessible at all times, reducing the need to return to the home screen after generating a plan.")
body(doc, "The visual design uses a neon-inspired cool color palette, compact planner controls, icon-led navigation and card-based itinerary output. This makes the application feel modern while still remaining usable for repeated planning tasks.")

heading(doc, "VIII. TESTING AND OBSERVATIONS")
body(doc, "The project includes Django tests for successful plan generation, long-trip validation, named halt creation, varied route leg distances and the coordinate-based weather endpoint. Browser verification was also performed on a Bangalore-to-Manali long-route scenario, where the system generated fourteen varied route segments and weather cards for all generated stops.")
body(doc, "The main observed improvement over the first prototype is that RoadNova no longer completes long trips on day one with buffer days afterward. Instead, it distributes the route across the selected date range and treats each day as a real driving stage.")

heading(doc, "IX. LIMITATIONS AND FUTURE WORK")
body(doc, "RoadNova is a working academic prototype and not yet a certified navigation or emergency system. Toll values are currently estimates, and hotel or restaurant suggestions depend on curated data and available OSM search results. Weather accuracy depends on the external forecast provider and selected forecast range.")
body(doc, "Future work can include live traffic, verified hotel inventory, user accounts, saved collaborative trips, offline emergency sharing, toll API integration, more Indian cities, route risk scoring and a self-hosted routing backend for production reliability.")

heading(doc, "X. CONCLUSION")
body(doc, "RoadNova demonstrates how a Django web application can combine itinerary generation, open maps, route geometry, weather forecasts, cost planning, SOS support and shared expense splitting into one India-focused road-trip planner. By enforcing feasible daily distances and using named intermediate halts, the system shifts from generic tourism listing to realistic road journey planning.")

heading(doc, "ACKNOWLEDGMENT")
body(doc, "The authors acknowledge the Department of Computer Science and Engineering, Dayananda Sagar Academy of Technology and Management, Bangalore, India, for supporting the project environment and academic guidance.")

heading(doc, "REFERENCES")
refs = [
    "P. Vansteenwegen and W. Souffriau, \"Trip planning functionalities: State of the art and future,\" Information Technology & Tourism, vol. 12, no. 4, pp. 305-315, 2011, doi: 10.3727/109830511X13049763021853.",
    "W. Souffriau, P. Vansteenwegen, J. Vertommen, G. Vanden Berghe and D. Van Oudheusden, \"A personalized tourist trip design algorithm for mobile tourist guides,\" Applied Artificial Intelligence, vol. 22, no. 10, pp. 964-985, 2008, doi: 10.1080/08839510802379626.",
    "T. Adamo, L. Colizzi, G. Dimauro, G. Ghiani and E. Guerriero, \"A multimodal tourist trip planner integrating road and pedestrian networks,\" Expert Systems with Applications, 2024, doi: 10.1016/j.eswa.2023.121457.",
    "M. Haklay and P. Weber, \"OpenStreetMap: User-generated street maps,\" IEEE Pervasive Computing, vol. 7, no. 4, pp. 12-18, 2008, doi: 10.1109/MPRV.2008.80.",
    "S. Huber and C. Rust, \"Calculate travel time and distance with OpenStreetMap data using the Open Source Routing Machine (OSRM),\" The Stata Journal, vol. 16, no. 2, pp. 416-423, 2016, doi: 10.1177/1536867X1601600209.",
    "T. Giraud, \"osrm: Interface between R and the OpenStreetMap-based routing service OSRM,\" Journal of Open Source Software, vol. 7, no. 78, p. 4574, 2022, doi: 10.21105/joss.04574.",
    "W. H. K. Lam, H. Shao and A. Sumalee, \"Modeling impacts of adverse weather conditions on a road network with uncertainties in demand and supply,\" Transportation Research Part B: Methodological, vol. 42, no. 10, pp. 890-910, 2008, doi: 10.1016/j.trb.2008.02.004.",
    "M. Zhang, X. Hu and J. Liao, \"Path optimization in dynamic adverse weathers,\" Journal of Risk Analysis and Crisis Response, vol. 6, pp. 197-205, 2016, doi: 10.2991/jrarc.2016.6.4.4.",
    "A. Fielbaum, X. Bai and J. Alonso-Mora, \"How to split the costs and charge the travellers sharing a ride? Aligning system's optimum with users' equilibrium,\" European Journal of Operational Research, vol. 301, no. 3, pp. 956-973, 2022, doi: 10.1016/j.ejor.2021.11.027.",
    "R. H. Howlader, \"Settling debts efficiently: Zero-sum set packing,\" Harvard University DASH, 2017.",
    "Open-Meteo, \"Weather Forecast API documentation,\" 2026. [Online]. Available: https://open-meteo.com/",
    "Project OSRM, \"Open Source Routing Machine backend,\" 2026. [Online]. Available: https://github.com/Project-OSRM/osrm-backend",
]
for i, ref in enumerate(refs, 1):
    add_reference(doc, i, ref)

doc.save(OUT)
print(OUT)
